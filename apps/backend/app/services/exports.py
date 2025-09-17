"""Submission export helpers."""
from __future__ import annotations

import csv
import io
from dataclasses import dataclass

from docx import Document
from fastapi import HTTPException

from ..models import Submission


@dataclass
class ExportPayload:
    filename: str
    media_type: str
    content: bytes


class ExportService:
    """Provide different export formats for submissions."""

    def __init__(self, submission: Submission) -> None:
        self._submission = submission

    def to_csv(self) -> ExportPayload:
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        writer.writerow(["Field", "Value"])
        writer.writerow(["Title", self._submission.title])
        writer.writerow(["Status", self._submission.status])
        writer.writerow(["Tone", self._submission.tone or ""]) 
        writer.writerow(["Audience", self._submission.audience or ""])
        writer.writerow(["Channel", self._submission.channel or ""])
        writer.writerow(["English Source", self._submission.source_text])
        writer.writerow(["Thai Draft", self._submission.thai_draft])
        writer.writerow(["Thai Final", self._submission.thai_final or ""])
        writer.writerow(["Reviewer Notes", self._submission.reviewer_notes or ""])
        writer.writerow(["Warnings", " | ".join(self._submission.warnings or [])])
        writer.writerow(["Glossary", " | ".join(self._submission.glossary_terms or [])])
        writer.writerow(["Provider", self._submission.provider_name or ""])
        writer.writerow(["Tokens", self._submission.usage_tokens or 0])
        writer.writerow(["Cost (USD)", self._submission.cost_usd or 0.0])

        return ExportPayload(
            filename=f"submission-{self._submission.id}.csv",
            media_type="text/csv",
            content=buffer.getvalue().encode("utf-8-sig"),
        )

    def to_docx(self) -> ExportPayload:
        document = Document()
        document.add_heading(self._submission.title, level=1)
        document.add_paragraph(f"Status: {self._submission.status}")
        document.add_paragraph(f"Tone: {self._submission.tone or 'default'}")
        document.add_paragraph(f"Audience: {self._submission.audience or 'unspecified'}")
        document.add_paragraph(f"Channel: {self._submission.channel or 'unspecified'}")

        document.add_heading("English Source", level=2)
        document.add_paragraph(self._submission.source_text)

        document.add_heading("Generated Draft", level=2)
        document.add_paragraph(self._submission.thai_draft)

        document.add_heading("Final Copy", level=2)
        document.add_paragraph(self._submission.thai_final or "Pending edits")

        if self._submission.reviewer_notes:
            document.add_heading("Reviewer Notes", level=2)
            document.add_paragraph(self._submission.reviewer_notes)

        if self._submission.warnings:
            document.add_heading("Warnings", level=2)
            for warning in self._submission.warnings:
                document.add_paragraph(warning, style="List Bullet")

        buffer = io.BytesIO()
        document.save(buffer)

        return ExportPayload(
            filename=f"submission-{self._submission.id}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            content=buffer.getvalue(),
        )

    def to_social(self) -> ExportPayload:
        # Provide a simple plaintext caption with CTA and notes for social scheduling
        text = io.StringIO()
        text.write(f"Title: {self._submission.title}\n")
        text.write(f"Status: {self._submission.status}\n")
        text.write("--- English Source ---\n")
        text.write(self._submission.source_text + "\n\n")
        text.write("--- Thai Final ---\n")
        thai_final = self._submission.thai_final or self._submission.thai_draft
        text.write(thai_final + "\n")
        if self._submission.warnings:
            text.write("\nWarnings:\n")
            for warning in self._submission.warnings:
                text.write(f"- {warning}\n")

        return ExportPayload(
            filename=f"submission-{self._submission.id}.txt",
            media_type="text/plain",
            content=text.getvalue().encode("utf-8"),
        )

    def generate(self, format: str) -> ExportPayload:
        format = format.lower()
        if format == "csv":
            return self.to_csv()
        if format == "docx":
            return self.to_docx()
        if format == "social":
            return self.to_social()
        raise HTTPException(status_code=400, detail="Unsupported export format")
